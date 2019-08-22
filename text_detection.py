
from imutils.object_detection import non_max_suppression
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
import numpy as np
import docx
import argparse
import time
import cv2
import pytesseract
import  array as arr


ap = argparse.ArgumentParser()
ap.add_argument("-i", "--image", type=str,
                help="Girişte bulunan görüntünün yolu")
ap.add_argument("-east", "--east", type=str,
                help="EAST metin detektörüne giriş yolu")
ap.add_argument("-c", "--min-confidence", type=float, default=0.5,
                help="Bir bölgeyi denetlemek için gereken minumum olasılık")
ap.add_argument("-w", "--width", type=int, default=320,
                help="Yeni resmin genişliği, minumum 320 olmalıdır.")
ap.add_argument("-e", "--height", type=int, default=320,
                help="Yeniden resimin yüksekliği, minimum 320 olmalıdır.")
args = vars(ap.parse_args())  # Gelen argümanı değişkene aktarıyoruz.
# "--abc" optional şekilde olmasını sağlar "-a" ise onun kısalmasıdır.
# -h şeklinde argüman gönderirsek, argüman hakkında yardım masajı gösterecektir.


image = cv2.imread(args["image"])
orig = image.copy()
(H, W) = image.shape[:2]  # Görüntü boyutlarını alır (satır,sütun)

(newW, newH) = (args["width"], args["height"])  # Yeni genişlik ve yüksekliği ayarlar.
rW = W / float(newW)  # Değişiklik oranı belirleme.
rH = H / float(newH)

# Yeni boyutlara göre resmi şekillendirir.
image = cv2.resize(image, (newW, newH))
(H2, W2) = image.shape[:2]

layerNames = [
    "feature_fusion/Conv_7/Sigmoid",  # Çıktı olasılıkları (...)
    "feature_fusion/concat_3"  # Metni çerçeveleyen kutu kordinatları (...)
            ]

# Önceden eğitilmiş EAST metin dedektörünü yükleme mesajı.
print("[INFO] loading EAST ext detector...")
net = cv2.dnn.readNet(args["east"])

# Görüntü oluşturur. (...)
blob = cv2.dnn.blobFromImage(image, 1.0, (W2, H2),
                             (123.68, 116.78, 103.94), swapRB=True, crop=False)
start = time.time()
net.setInput(blob)
(scores, geometry) = net.forward(layerNames)
end = time.time()

# Tahmini süre bilgisi
print("[INFO] text detection took {:.6f} seconds".format(end - start))  # (...)

(numRows, numCols) = scores.shape[2:4]
rects = []  # Yazıyı çevreleyen kutunun kordinatlarını saklar.
confidences = []  # Yazıyı çevreleyen kutuların olasılığını hesaplar.


# Satır sayısı üzerine döngü
for y in range(0, numRows):
    # Puanları (olasılıkları), ardından geometrikleri çıkar
    # Potansiyel sınırlayıcı kutu koordinatlarını türetmek için kullanılan veri.
    # Metni çevrele

    scoresData = scores[0, 0, y]
    xData0 = geometry[0, 0, y]
    xData1 = geometry[0, 1, y]
    xData2 = geometry[0, 2, y]
    xData3 = geometry[0, 3, y]
    anglesData = geometry[0, 4, y] # Köşe bilgileri.


    # Sütun sayısını döngüye koyarız.
    for x in range(0, numCols):

        # Eğer kelimemizin olasılık değeri yeterli değilse görmezden geleceğiz.
        if scoresData[x] < args["min_confidence"]:
            continue

        # Ortaya çıkan özellik ofset faktörünü hesaplar.
        # Giriş görüntüsünden 4x daha küçük.
        (offsetX, offsetY) = (x * 4.0, y * 4.0)

        # Tahmin için dönme açısının farkını bulur.
        # Sin ve cos değerini hesaplar.
        angle = anglesData[x]
        cos = np.cos(angle)
        sin = np.sin(angle)

       # Genişliğini ve yüksekliğini türetmek için geometri hacmini kullanın
       # Sınırlayıcı kutu
        h = xData0[x] + xData2[x]
        w = xData1[x] + xData3[x]

        # Metin tahmini sınırlama kutusu için hem başlangıç hem de bitiş (x, y) koordinatlarını hesaplar
        endX = int(offsetX + (cos * xData1[x]) + (sin * xData2[x]))
        endY = int(offsetY - (sin * xData1[x]) + (cos * xData2[x]))
        startX = int(endX - w)
        startY = int(endY - h)

        # Sınırlama kutusu koordinatlarını ve olasılık puanını ilgili listemize ekleyin
        rects.append((startX, startY, endX, endY))
        confidences.append(scoresData[x])

        # Zayıf, üst üste binen sınırlayıcı kutuları gizlemek için maksimum olmayan baskı uygulanır.
        boxes = non_max_suppression(np.array(rects), probs=confidences)


Kelime = []
Genis = []
Uzun = []
Xkoordinat = []
Ykoordinat = []

# Metni çerçevesi üzerine döngü.
for (startX, startY, endX, endY) in boxes:
        # Sınırlayıcı kutu koordinatlarının ilgili oranlara göre ölçeklendirilmesi
        startX = int(startX * rW)
        startY = int(startY * rH)
        endX = int(endX * rW)
        endY = int(endY * rH)
        genislik= ( endX - startX )
        uzunluk= (endY - startY)
        # print(confidences)
        #helv36 = tkFont.Font(family = "Helvetica",size = 36,weight = "bold")




        # Resmin üzerine kutu çizilir
        cv2.rectangle(orig, (startX, startY), (endX, endY), (0, 255, 0), 2)


        # Çerçevenin içindeki yazının okunmasını sağlıyor.
        roi = orig[startY:endY, startX:endX]
        config = ("-l eng --oem 1 --psm 7")
        result = pytesseract.image_to_string(roi, config=config)


        # Çıktıları listeye atıyor.
        Kelime.append(result)
        Xkoordinat.append(startX)
        Ykoordinat.append(startY)
        Genis.append(genislik)
        Uzun.append(uzunluk)

        # Güven değerine göre sıralanmış olan kelimeleri x ve y koordinatlarına göre yeniden sıralar.
        # Buble Sort algoritması kullanır.
for i in range(0, len(Ykoordinat) - 1):
    for j in range(0, len(Ykoordinat) - 1 - i):
        for i in range(0, len(Xkoordinat) - 1):
            for j in range(0, len(Xkoordinat) - 1 - i):
                if Ykoordinat[j] > Ykoordinat[j + 1]: #and Ykoordinat[i] - Ykoordinat[i+1] < Uzun[i]:
                    Ykoordinat[j], Ykoordinat[j + 1] = Ykoordinat[j + 1], Ykoordinat[j]
                    Xkoordinat[j], Xkoordinat[j + 1] = Xkoordinat[j + 1], Xkoordinat[j]
                    Kelime[j], Kelime[j + 1] = Kelime[j + 1], Kelime[j]
                    Genis[j], Genis[j + 1] = Genis[j + 1], Genis[j]
                    Uzun[j], Uzun[j + 1] = Uzun[j + 1], Uzun[j]

# Yazı fontu
document = Document()
font = document.styles['Normal'].font
font.name = 'Arial'


# Yazdırma düzenini ayarlar
sections = document.sections
for section in sections:
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.75)
    section.right_margin = Cm(1.75)

# Yazdırma boyutunu ayarlar
p = document.add_paragraph()
for i in range(0, len(Kelime)-1):
    font.size = Pt((Uzun[i]*500)/H)
    p.add_run(Kelime[i]+ " ")
    # Alt satıra geçmesini sağlıyor.
    if ( (Ykoordinat[i+1]-Ykoordinat[i]) >= Uzun[i]):
        p = document.add_paragraph()
p.add_run(Kelime[i+1]+ " ")
document.save(r'Document.docx')


print(Kelime)
print(Uzun)
# Çıkış resmini göster.
cv2.imshow("Text Detection", orig)
cv2.waitKey(0)






